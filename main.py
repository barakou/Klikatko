import datetime
import subprocess
import numpy as np
from PyQt5 import QtGui
from docx import Document
from docx.shared import Pt, Inches
from PyQt5.QtWidgets import QApplication, QDialog, QLabel, QLineEdit, QVBoxLayout, QRadioButton, QPushButton, QCheckBox, \
    QFileDialog, QComboBox, QGroupBox

#Set czech datetime
# Slovník pro překlad anglických názvů měsíců do češtiny
mesice_cz = {
    'January': 'leden',
    'February': 'únor',
    'March': 'březen',
    'April': 'duben',
    'May': 'květen',
    'June': 'červen',
    'July': 'červenec',
    'August': 'srpen',
    'September': 'září',
    'October': 'říjen',
    'November': 'listopad',
    'December': 'prosinec'
}
date = datetime.date.today()
year = date.year
month = mesice_cz.get(date.strftime('%B'), 'Neznámý')

sloucena_dokumentace = []
stavebni_povoleni = []
uzemni_rizeni = []

# Create a QApplication instance
app = QApplication([])

# Create a QDialog window for the input and radio buttons
input_dialog = QDialog()
input_dialog.setWindowTitle("Klikátko™")
input_dialog.setWindowIcon(QtGui.QIcon('logo.png'))

# Create a QVBoxLayout to hold the widgets
layout = QVBoxLayout(input_dialog)

# Document Name Input
name_label = QLabel("Zadejte název dokumentu:")
layout.addWidget(name_label)
document_name = QLineEdit()
layout.addWidget(document_name)

# Contract Name Input
contract_label = QLabel("Zadejte číslo zakázky:")
layout.addWidget(contract_label)
contract_name = QLineEdit()
layout.addWidget(contract_name)

# Radio Buttons
option_label = QLabel("Typ průvodní zprávy:")
layout.addWidget(option_label)

option1_radiobutton = QRadioButton("NEUMIM Dokumentace pro územní řízení")
layout.addWidget(option1_radiobutton)

option2_radiobutton = QRadioButton("NEUMIM Dokumentace pro stavební povolení")
layout.addWidget(option2_radiobutton)

option3_radiobutton = QRadioButton("Sloučená dokumentace pro stavební a územní řízení")
layout.addWidget(option3_radiobutton)

# Connect the radiobuttons to update the selected_option variable
def update_selected_option():
    global selected_option
    if option1_radiobutton.isChecked():
        selected_option = "Průvodní a souhrnná zpráva"
    elif option2_radiobutton.isChecked():
        selected_option = "Stavební řízení - neimplementováno"
    elif option3_radiobutton.isChecked():
        selected_option = "Stavební řízení - neimplementováno"

option1_radiobutton.toggled.connect(update_selected_option)
option2_radiobutton.toggled.connect(update_selected_option)
option3_radiobutton.toggled.connect(update_selected_option)


image_button = QPushButton("Vybrat obrázek")
image_button.setEnabled(True)
image_label = QLabel()
global image_path
layout.addWidget(image_label)
layout.addWidget(image_button)

# Function to handle image button click
def select_image():
    global image_path
    file_dialog = QFileDialog()
    file_dialog.setWindowTitle("Vyberte obrázek")
    file_dialog.setFileMode(QFileDialog.ExistingFile)
    if file_dialog.exec_() == QDialog.Accepted:
        selected_files = file_dialog.selectedFiles()
        if selected_files:
            image_path= selected_files[0]
            # Do something with the selected image path
            # For example, you can store it in a variable or display it in the UI
        else:
            # No image selected

            image_label.setText("<font color='red'>Nebyl vybrán žádný obrázek</font>")
            image_path = ""

image_button.clicked.connect(select_image)



# Button to continue
continue_button = QPushButton("Pokračovat")
continue_button.clicked.connect(input_dialog.accept)
layout.addWidget(continue_button)

# Set focus on the document name entry field
document_name.setFocus()

# Execute the QDialog window
if input_dialog.exec_() == QDialog.Accepted:
    document_name_text = document_name.text()
    contract_name_text = contract_name.text()

    if document_name_text:
        checkbox_dialog = QDialog()
        checkbox_dialog.setWindowTitle("Klikátko™ : Nastavení dokumentu")
        checkbox_dialog.setWindowIcon(QtGui.QIcon('logo.png'))
        checkbox_dialog.resize(500, 300)

        # Create a QVBoxLayout to hold the checkboxes and close button
        checkbox_layout = QVBoxLayout(checkbox_dialog)

        # Checkbox options
        checkbox_label = QLabel("Sekce dokumentu:")
        checkbox_layout.addWidget(checkbox_label)

        def create_subsection(paragraph_name, is_enabled, section):
            option_checkbox = QCheckBox(f"{paragraph_name}")
            option_combobox = QComboBox()
            option_checkbox.setChecked(is_enabled)
            option_combobox.addItem("Netýká se")
            option_combobox.addItem("Doplnit")
            option_combobox.setEnabled(is_enabled)
            section.addWidget(option_checkbox)
            section.addWidget(option_combobox)
            return option_checkbox, option_combobox


        # Create the first section with checkboxes and dropdown menus
        section1_groupbox = QGroupBox("Sekce A")
        section1_layout = QVBoxLayout()
        section1_groupbox.setLayout(section1_layout)

        option1_checkbox, option1_combobox = create_subsection("A1. Identifikační údaje", True, section1_layout)
        option2_checkbox, option2_combobox = create_subsection("A2. Členění stavby", True, section1_layout)
        option3_checkbox, option3_combobox = create_subsection("A3. Seznam vstupních podkladů", True, section1_layout)

        checkbox_layout.addWidget(section1_groupbox)

        # Create the second section with checkboxes and dropdown menus
        section2_groupbox = QGroupBox("B Souhrnná zpráva")
        section2_layout = QVBoxLayout()
        section2_groupbox.setLayout(section2_layout)

        option4_checkbox, option4_combobox = create_subsection("B1. Popis území stavby", True, section2_layout)
        option5_checkbox, option5_combobox = create_subsection("B2. Celkový popis stavby", True, section2_layout)
        option6_checkbox, option6_combobox = create_subsection("B3. Připojení stavby na technickou infrastrukturu", False, section2_layout)
        option7_checkbox, option7_combobox = create_subsection("B4", False, section2_layout)
        option8_checkbox, option8_combobox = create_subsection("B5", False, section2_layout)
        option9_checkbox, option9_combobox = create_subsection("B6", False, section2_layout)
        option10_checkbox, option10_combobox = create_subsection("B7", False, section2_layout)
        option11_checkbox, option11_combobox = create_subsection("B8", False, section2_layout)
        option12_checkbox, option12_combobox = create_subsection("B9", False, section2_layout)

        checkbox_layout.addWidget(section2_groupbox)


        # Function to update the enabled state of the dropdown menus
        def update_dropdown_state():
            option1_combobox.setEnabled(option1_checkbox.isChecked())
            option2_combobox.setEnabled(option2_checkbox.isChecked())
            option3_combobox.setEnabled(option2_checkbox.isChecked())
            option4_combobox.setEnabled(option4_checkbox.isChecked())
            option5_combobox.setEnabled(option5_checkbox.isChecked())
            option6_combobox.setEnabled(option6_checkbox.isChecked())
            option7_combobox.setEnabled(option7_checkbox.isChecked())
            option8_combobox.setEnabled(option8_checkbox.isChecked())
            option9_combobox.setEnabled(option9_checkbox.isChecked())
            option10_combobox.setEnabled(option10_checkbox.isChecked())
            option11_combobox.setEnabled(option11_checkbox.isChecked())
            option12_combobox.setEnabled(option12_checkbox.isChecked())


        # Connect the checkbox state changed signals to the update function
        option1_checkbox.stateChanged.connect(update_dropdown_state)
        option2_checkbox.stateChanged.connect(update_dropdown_state)
        option3_checkbox.stateChanged.connect(update_dropdown_state)
        option4_checkbox.stateChanged.connect(update_dropdown_state)
        option5_checkbox.stateChanged.connect(update_dropdown_state)
        option6_checkbox.stateChanged.connect(update_dropdown_state)
        option7_checkbox.stateChanged.connect(update_dropdown_state)
        option8_checkbox.stateChanged.connect(update_dropdown_state)
        option9_checkbox.stateChanged.connect(update_dropdown_state)
        option10_checkbox.stateChanged.connect(update_dropdown_state)
        option11_checkbox.stateChanged.connect(update_dropdown_state)
        option12_checkbox.stateChanged.connect(update_dropdown_state)

        # Button to save and close
        def save_clicked():
            # Retrieve checkbox selections
            selected_options = []
            if option1_checkbox.isChecked():
                selected_options.append("Option 1")
            if option2_checkbox.isChecked():
                selected_options.append("Option 2")

            # Perform further processing with the selected options
            def generate_paragraph(paragraphs):
                for index, text in enumerate(paragraphs, start=1):
                    p = document.add_paragraph(style='AlphabeticBullet')
                    run = p.add_run(f"{chr(ord('a') + index - 1)}) ")  # Add the alphabetic bullet
                    run.bold = True
                    p.add_run(text)

            # Create and format the document
            document = Document()

            # document.add_section()
            # Add the first section
            section1 = document.sections[0]

            # Set different header and footer for the first page
            header_first_page = section1.header
            header_first_page.is_linked_to_previous = True  # Remove the link to previous header

            footer_first_page = section1.footer
            footer_first_page.is_linked_to_previous = False
            # Add content to the footer of the first page
            f_paragraph_first_page = footer_first_page.paragraphs[0]
            f_paragraph_first_page.text = f"{month} {year}" "\t\tVypracoval: Ing. Koudelka"
            f_paragraph_first_page.style = document.styles["Footer"]

            for paragraph in header_first_page.paragraphs:
                paragraph.text = ""

            p = document.add_paragraph(f"{document_name_text}.")
            document.add_paragraph(
                "Jednostupňová dokumentace pro sloučené územní a stavební řízení dle vyhlášky 583/2020 Sb.")
            document.add_paragraph("Číslo zakázky: "f"{contract_name_text}")

            if (image_path):
                image = document.add_picture(image_path, width=Inches(7.0))
            else:
                image = document.add_picture("image-filename.png", width=Inches(7.0))

            image_width = Inches(6.5)
            image_height = image.height

            # Get the page width and height
            page_width = document.sections[0].page_width
            page_height = document.sections[0].page_height

            # Calculate the position to center the image
            left = (page_width - image_width) / 2
            top = (page_height - image_height) / 2

            # Set the image position
            image.left = Inches(left)
            image.top = Inches(top)

            document.add_paragraph("Průvodní a souhrnná zpráva")
            document.add_paragraph("Příloha A-B")

            # Add a section break and start a new section for the second page
            document.add_section()

            # Set different header and footer for the second page onwards
            section2 = document.sections[1]
            header_second_page = section2.header
            paragraph_header_second_page = header_second_page.paragraphs[0]
            paragraph_header_second_page.text = f"{document_name_text} \t\tPrůvodní a souhrnná zpráva"
            paragraph_header_second_page.style = document.styles["Header"]

            footer_second_page = section2.footer
            footer_second_page.is_linked_to_previous = False

            # Clear the content of the footer for the second page onwards
            for paragraph in footer_second_page.paragraphs:
                for run in paragraph.runs:
                    run.clear()

            bullet_style = document.styles.add_style('AlphabeticBullet', 1)
            # bullet_style.base_style = document.styles['ListBullet']
            bullet_formatting = bullet_style.paragraph_format
            bullet_formatting.left_indent = Pt(18)

            document.add_heading("A. Průvodní zpráva", level=1)
            document.add_heading("A1. Identifikační údaje", level=2)
            # document.add_paragraph = "Stavba: \t\t" f"{document_name_text}"
            document.add_heading("A2. Členění stavby na objekty a technická a technologická zařízení", level=2)
            document.add_heading("A3. Seznam vstupních podkladů", level=2)
            document.add_heading("B. Souhrnná zpráva:", level=1)
            document.add_heading("B1. Popis území stavby", level=2)
            document.add_heading("B2. Celkový popis stavby", level=2)
            document.add_heading("B2.1 Základní charakteristika stavby a její používání", level=3)
            document.add_heading("B2.2 Celkové urbanistické a architektonické řešení", level=3)
            document.add_heading("B2.3 Celkové stavebně technické řešení", level=3)
            document.add_heading("B2.4 Bezbariérové užívání stavby", level=3)
            document.add_heading("B2.5 Bezpečnost při užívání stavby", level=3)
            document.add_heading("B2.6 Základní technický popis stavebních objektů", level=3)
            document.add_heading("B2.7 Základní popis technických a technologických objektů", level=3)
            document.add_heading("B2.8 Zásady požárně bezpečnostního řešení", level=3)
            document.add_heading("B2.9 Úspora energie a tepelná ochrana", level=3)
            document.add_heading("B2.10 Hygienické požadavky na stavbu", level=3)
            document.add_heading("B2.11 Zásady ochrany stavby před negativními účinky vnějšího prostředí", level=3)
            document.add_heading("B3. Připojení stavby na technickou infrastrukturu", level=2)
            paragraphs = [
                "napojovací místa technické infrastruktury",
                "připojovací rozměry, výkonové kapacity a délky"
            ]
            generate_paragraph(paragraphs)
            document.add_heading("B4. Dopravní řešení a základní údaje o provozu, provozní a dopravní technologie",
                                 level=2)
            paragraphs = [
                "popis dopravního řešení",
                "bezbariérové opatření pro přístupnost",
                "doprava v klidu",
                "pěší a cyklistické stezky"
            ]
            generate_paragraph(paragraphs)
            document.add_heading("B5. Řešení vegetace a sovisejících terénních úprav", level=2)
            paragraphs = [
                "terénní úpravy",
                "použité vegetační prvky",
                "biotechnická, protierozní opatření"
            ]
            generate_paragraph(paragraphs)
            document.add_heading("B6. Popis vlivů stavby na životní prostředí a jeho ochrana", level=2)
            paragraphs = [
                "vliv na životní prostředí, ovzduší, hluk, voda, odpady, půda a horninové prostředí",
                "vliv na přírodu a krajinu,krajinný ráz, přírodní parky, dřeviny, památné stromy",
                "vliv na  soustavu chráněných území Natura 2000",
                "způsob zohlednění podmínek závazného stanoviska posouzení vlivu záměru na životní prostředí, je-li podkladem",
                "popis souladu záměru s oznámením záměru dle zákona o posuzování vlivu záměru na životní prostředí, je-li podkladem",
                "navrhovaná ochranná a bezpečnostní pásma, rozsah omezení a podmínky ochrany podle jiných právních předpisů"
            ]
            generate_paragraph(paragraphs)
            document.add_heading("B7. Ochrana obyvatelstva", level=2)
            paragraphs = [
                "opatření vyplývající z požadavků civilní ochrany",
                "prevence závažných havárií"
            ]
            generate_paragraph(paragraphs)
            document.add_heading("B8. Zásady organizace výstavby", level=2)
            paragraphs = [
                "napojení staveniště na stávající dopravní a technickou infrastrukturu",
                "přístup na stavbu po dobu výstavby, popřípadě přístupové trasy",
                "ochrana okolí staveniště a požadavky na související asanace, demolice, kácení dřevin",
                "maximální dočasné a trvalé zábory pro staveniště",
                "požadavky na bezbariérové obchozí trasy",
                "základní bilance zemních prací, požadavky na přísun nebo deponie zemin",
                "návrh postupu výstavby(časový plán, harmonogramy, etapizace, výluky apod.)",
                "požadavky na postupné uvádění stavby do provozu (užívání), požadavky na průběh a způsob přípravy a realizace výstavby",
                "návrh objízdných tras pro automobily, veřejnou dopravu, cyklisty a pěší, včetně průchodů pěších staveništěm v jednotlivých stavebních etapách (DIO)"
            ]
            generate_paragraph(paragraphs)

            document.add_heading("B9. Celkové vodohospodářské řešení", level=2)

            document.add_page_break()

            # Launch Word and open the document
            subprocess.Popen(["start", "winword", "A - B průvodní a souhrnná zpráva.docx"], shell=True)

            # Save the document
            document.save("A - B průvodní a souhrnná zpráva.docx")

            # Close the checkbox window
            checkbox_dialog.accept()

        save_button = QPushButton("Uložit")
        save_button.clicked.connect(save_clicked)
        checkbox_layout.addWidget(save_button)

        # Button to close the application
        close_button = QPushButton("Zavřít")
        close_button.clicked.connect(checkbox_dialog.reject)
        checkbox_layout.addWidget(close_button)

        # Execute the checkbox QDialog window
        checkbox_dialog.exec_()


